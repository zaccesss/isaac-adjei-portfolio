#!/usr/bin/env python3
"""
Generate clean DOCX files from HTML CVs without JavaScript.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_html_for_docx(html_content):
    """Remove JavaScript and clean HTML for docx conversion."""
    # Remove script tags and their content
    html_content = re.sub(r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>', '', html_content, flags=re.DOTALL)
    # Remove onload attributes
    html_content = re.sub(r'onload="[^"]*"', '', html_content)
    # Remove onclick attributes  
    html_content = re.sub(r'onclick="[^"]*"', '', html_content)
    return html_content

def html_to_docx(html_file, docx_file):
    """Convert HTML CV to DOCX."""
    with open(html_file, 'r', encoding='utf-8') as f:
        html = f.read()
    
    # Clean the HTML
    html = clean_html_for_docx(html)
    
    # Create a new Document
    doc = Document()
    
    # Extract text content using regex
    # Header - Name
    name_match = re.search(r'<div class="header-name">(.*?)</div>', html)
    if name_match:
        name = name_match.group(1)
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_match = re.search(r'<div class="header-contact">(.*?)</div>', html, re.DOTALL)
    if contact_match:
        contact_html = contact_match.group(1)
        # Remove SVG tags
        contact_text = re.sub(r'<svg[^>]*>.*?</svg>', ' ', contact_html, flags=re.DOTALL)
        # Remove remaining HTML tags
        contact_text = re.sub(r'<[^>]+>', ' ', contact_text)
        # Clean up whitespace
        contact_text = re.sub(r'\s+', ' ', contact_text).strip()
        contact_text = contact_text.replace(' | ', ' | ')
        
        p = doc.add_paragraph()
        run = p.add_run(contact_text)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Profile
    profile_match = re.search(r'<div class="profile">(.*?)</div>', html, re.DOTALL)
    if profile_match:
        profile_text = re.sub(r'<[^>]+>', '', profile_match.group(1))
        profile_text = re.sub(r'\s+', ' ', profile_text).strip()
        
        doc.add_heading('Profile', level=1)
        p = doc.add_paragraph(profile_text)
        p.paragraph_format.space_after = Pt(6)
    
    # Education
    edu_section = re.search(r'<!-- EDUCATION -->(.*?)<!-- SKILLS -->', html, re.DOTALL)
    if edu_section:
        doc.add_heading('Education', level=1)
        edu_html = edu_section.group(1)
        
        # Extract education entries
        entries = re.findall(r'<div class="edu-entry">(.*?)</div>\s*</div>', edu_html, re.DOTALL)
        for entry in entries:
            # Institution and date
            inst_match = re.search(r'<span class="edu-left">(.*?)</span>.*?<span class="edu-right">(.*?)</span>', entry, re.DOTALL)
            if inst_match:
                institution = re.sub(r'<[^>]+>', '', inst_match.group(1))
                date = re.sub(r'<[^>]+>', '', inst_match.group(2))
                
                p = doc.add_paragraph()
                run = p.add_run(institution)
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(' ' + date).italic = True
            
            # Degree and award
            degree_match = re.search(r'<div class="edu-sub">(.*?)</div>', entry, re.DOTALL)
            if degree_match:
                degree = re.sub(r'<[^>]+>', '', degree_match.group(1))
                p = doc.add_paragraph(degree)
                p.paragraph_format.left_indent = Inches(0.25)
    
    # Skills
    skills_section = re.search(r'<!-- SKILLS -->(.*?)<!-- PROJECTS -->', html, re.DOTALL)
    if skills_section:
        doc.add_heading('Skills (Professional and Technical)', level=1)
        skills_html = skills_section.group(1)
        
        # Extract skill categories
        categories = re.findall(r'<p class="skill-category"><span class="skill-label">(.*?):</span>(.*?)</p>', skills_html, re.DOTALL)
        for label, skills in categories:
            label = re.sub(r'<[^>]+>', '', label)
            skills = re.sub(r'<[^>]+>', '', skills).strip()
            
            p = doc.add_paragraph()
            run = p.add_run(label + ': ')
            run.bold = True
            p.add_run(skills)
    
    # Projects
    projects_section = re.search(r'<!-- PROJECTS -->(.*?)<!-- EXPERIENCE -->', html, re.DOTALL)
    if projects_section:
        doc.add_heading('Projects', level=1)
        projects_html = projects_section.group(1)
        
        # Extract project entries
        proj_entries = re.findall(r'<div class="proj-entry">(.*?)</div>\s*</div>', projects_html, re.DOTALL)
        for entry in proj_entries:
            # Project name and date
            name_match = re.search(r'<span class="proj-name">(.*?)</span>.*?<span class="proj-date">(.*?)</span>', entry, re.DOTALL)
            if name_match:
                name = re.sub(r'<[^>]+>', '', name_match.group(1))
                date = re.sub(r'<[^>]+>', '', name_match.group(2))
                
                p = doc.add_paragraph()
                run = p.add_run(name)
                run.bold = True
                p.add_run(' ' + date).italic = True
            
            # Tech stack
            tech_match = re.search(r'<div class="proj-tech">(.*?)</div>', entry, re.DOTALL)
            if tech_match:
                tech = re.sub(r'<[^>]+>', '', tech_match.group(1))
                p = doc.add_paragraph(tech)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)
            
            # Bullets
            bullets = re.findall(r'<li>(.*?)</li>', entry, re.DOTALL)
            for bullet in bullets:
                bullet_text = re.sub(r'<[^>]+>', '', bullet).strip()
                if bullet_text:
                    p = doc.add_paragraph('• ' + bullet_text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
    
    # Experience
    exp_section = re.search(r'<!-- EXPERIENCE -->(.*?)<!-- VOLUNTEERING -->', html, re.DOTALL)
    if exp_section:
        doc.add_heading('Experience', level=1)
        exp_html = exp_section.group(1)
        
        # Extract experience entries
        exp_entries = re.findall(r'<div class="exp-entry">(.*?)</div>\s*</div>', exp_html, re.DOTALL)
        for entry in exp_entries:
            # Company and date
            company_match = re.search(r'<span class="exp-company">(.*?)</span>.*?<span class="exp-date">(.*?)</span>', entry, re.DOTALL)
            if company_match:
                company = re.sub(r'<[^>]+>', '', company_match.group(1))
                date = re.sub(r'<[^>]+>', '', company_match.group(2))
                
                p = doc.add_paragraph()
                run = p.add_run(company)
                run.bold = True
                p.add_run(' ' + date).italic = True
            
            # Role and tags
            role_match = re.search(r'<div class="exp-role">(.*?)</div>', entry, re.DOTALL)
            if role_match:
                role = re.sub(r'<[^>]+>', '', role_match.group(1))
                p = doc.add_paragraph(role)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(3)
            
            # Bullets
            bullets = re.findall(r'<li>(.*?)</li>', entry, re.DOTALL)
            for bullet in bullets:
                bullet_text = re.sub(r'<[^>]+>', '', bullet).strip()
                if bullet_text:
                    p = doc.add_paragraph('• ' + bullet_text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
    
    # Volunteering
    vol_section = re.search(r'<!-- VOLUNTEERING -->(.*?)<!-- LANGUAGES -->', html, re.DOTALL)
    if vol_section:
        doc.add_heading('Volunteering', level=1)
        vol_html = vol_section.group(1)
        
        # Extract volunteering entries
        vol_entries = re.findall(r'<div class="vol-entry">(.*?)</div>\s*</div>', vol_html, re.DOTALL)
        for entry in vol_entries:
            # Org and date
            org_match = re.search(r'<span class="vol-org">(.*?)</span>.*?<span class="vol-date">(.*?)</span>', entry, re.DOTALL)
            if org_match:
                org = re.sub(r'<[^>]+>', '', org_match.group(1))
                date = re.sub(r'<[^>]+>', '', org_match.group(2))
                
                p = doc.add_paragraph()
                run = p.add_run(org)
                run.bold = True
                p.add_run(' ' + date).italic = True
            
            # Bullets
            bullets = re.findall(r'<li>(.*?)</li>', entry, re.DOTALL)
            for bullet in bullets:
                bullet_text = re.sub(r'<[^>]+>', '', bullet).strip()
                if bullet_text:
                    p = doc.add_paragraph('• ' + bullet_text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
    
    # Languages
    lang_section = re.search(r'<!-- LANGUAGES -->(.*?)</body>', html, re.DOTALL)
    if lang_section:
        doc.add_heading('Spoken Languages', level=1)
        lang_html = lang_section.group(1)
        
        # Extract language content
        lang_text = re.sub(r'<[^>]+>', '', lang_html)
        lang_text = re.sub(r'\s+', ' ', lang_text).strip()
        
        if lang_text:
            doc.add_paragraph(lang_text)
    
    # Set narrow margins for compact layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Save
    doc.save(docx_file)
    print(f"Created: {docx_file}")

def main():
    resume_dir = Path(__file__).parent.parent / 'public' / 'resume'
    
    # Generate for main CV
    html_file = resume_dir / 'cv.html'
    docx_file = resume_dir / 'Isaac_Adjei_CV.docx'
    html_to_docx(html_file, docx_file)
    
    # Generate for role-specific CVs
    roles = ['software', 'embedded', 'data', 'devops', 'quant', 'security']
    for role in roles:
        html_file = resume_dir / f'cv-{role}.html'
        docx_file = resume_dir / f'cv-{role}.docx'
        if html_file.exists():
            html_to_docx(html_file, docx_file)
        else:
            print(f"Skipping {role} - HTML file not found")

if __name__ == '__main__':
    main()
